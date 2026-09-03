#!/usr/bin/env python3
"""Small, conservative DOCX edit entrypoint for text replacement."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path


def fail(code: str, message: str) -> None:
    json.dump({"status": "error", "code": code, "message": message}, sys.stdout, ensure_ascii=False)
    sys.stdout.write("\n")
    raise SystemExit(2)


def main() -> None:
    parser = argparse.ArgumentParser(description="Replace text in a DOCX through the Hatch Documents Skill.")
    parser.add_argument("input", type=Path)
    parser.add_argument("--output", type=Path, required=True)
    parser.add_argument("--find", required=True)
    parser.add_argument("--replace", required=True)
    args = parser.parse_args()
    if not args.input.is_file():
        fail("input_not_found", f"DOCX does not exist: {args.input}")
    if args.input.resolve() == args.output.resolve():
        fail("invalid_argument", "--output must be different from the input; the source document is never overwritten")
    if args.input.suffix.lower() in {".docm", ".dotm"}:
        fail("unsupported_format", "Text replacement does not edit macro-enabled Word packages because python-docx cannot preserve VBA parts; use accept_changes.py for package-preserving revision acceptance")
    if not args.find:
        fail("invalid_argument", "--find must not be empty")
    try:
        from docx import Document
    except ImportError as exc:
        fail("dependency_unavailable", f"python-docx is required: {exc}")
    try:
        document = Document(str(args.input))
    except Exception as exc:
        fail("invalid_document", f"Could not open DOCX: {exc}")

    replacement_count = 0
    formatting_fallbacks = 0

    def replace_paragraph(paragraph) -> None:
        nonlocal replacement_count, formatting_fallbacks
        replaced_in_run = False
        for run in paragraph.runs:
            count = run.text.count(args.find)
            if count:
                run.text = run.text.replace(args.find, args.replace)
                replacement_count += count
                replaced_in_run = True
        if not replaced_in_run and args.find in paragraph.text:
            paragraph.text = paragraph.text.replace(args.find, args.replace)
            replacement_count += 1
            formatting_fallbacks += 1

    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for section in document.sections:
        paragraphs.extend(section.header.paragraphs)
        paragraphs.extend(section.footer.paragraphs)
    for paragraph in paragraphs:
        replace_paragraph(paragraph)

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(args.output))
    json.dump({
        "status": "ok",
        "operation": "replace",
        "input": str(args.input),
        "output": str(args.output),
        "replacements": replacement_count,
        "formatting_fallbacks": formatting_fallbacks,
        "warning": "Some replacements crossed run boundaries and may need visual review." if formatting_fallbacks else None
    }, sys.stdout, ensure_ascii=False)
    sys.stdout.write("\n")


if __name__ == "__main__":
    main()
